import customtkinter as ctk
from tkinter import messagebox
from PIL import Image
from database import Database
from models import Customer
from utils import format_money

FON_RANGI = "#0a0a0a"
SARIQ_RANG = "#f1c40f"

class MarketApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.db = Database()
        self.title("SHOHJAXON QURILISH MARKET")
        self.geometry("1150x750")
        self.configure(fg_color=FON_RANGI)

        # 1. BOSH SAHIFA
        ctk.CTkLabel(self, text="SHOHJAXON QURILISH MARKET", font=("Oswald", 45, "bold"), text_color=SARIQ_RANG).pack(pady=30)
        mid_cont = ctk.CTkFrame(self, fg_color="transparent")
        mid_cont.pack(fill="both", expand=True, padx=50)

        menu_fr = ctk.CTkFrame(mid_cont, fg_color="transparent")
        menu_fr.pack(side="left", fill="y", pady=20)

        # "Mijozlar ro'yxati" endi menyuni ochadi
        ctk.CTkButton(menu_fr, text="Mijozlar ro'yxati", width=300, height=65, font=("Arial", 19, "bold"), 
                      fg_color="transparent", border_width=2, border_color=SARIQ_RANG, text_color=SARIQ_RANG, 
                      command=self.open_customers_menu).pack(pady=15)
        
        # 32-qatordan boshlab almashtiring:
        ctk.CTkButton(menu_fr, text="Qarz qo'shish", width=300, height=65, font=("Arial", 19, "bold"), 
                      fg_color="transparent", border_width=2, border_color=SARIQ_RANG, text_color=SARIQ_RANG, 
                      command=self.open_add_debt_window).pack(pady=15)

        ctk.CTkButton(menu_fr, text="To'lov qilish", width=300, height=65, font=("Arial", 19, "bold"), 
                      fg_color="transparent", border_width=2, border_color=SARIQ_RANG, text_color=SARIQ_RANG,
                      command=self.open_payment_window).pack(pady=15)
        # Yangi qo'shilgan tugma
        ctk.CTkButton(menu_fr, text="Hisob-kitoblar", width=300, height=65, font=("Arial", 19, "bold"), 
                      fg_color="transparent", border_width=2, border_color=SARIQ_RANG, text_color=SARIQ_RANG,
                      command=lambda: self.open_calculations_window()).pack(pady=15)
        self.add_logo(mid_cont)
        
        stat_fr = ctk.CTkFrame(self, height=130, fg_color="transparent")
        stat_fr.pack(fill="x", side="bottom", pady=40, padx=50)
        self.lbl_debtors = self.create_stat_box(stat_fr, "Qarzdorlar soni: 0 ta", "left")
        self.lbl_total = self.create_stat_box(stat_fr, "Umumiy qarz: 0 so'm", "right")
        self.update_stats()

    def add_logo(self, parent):
        logo_fr = ctk.CTkFrame(parent, fg_color="transparent")
        logo_fr.pack(side="right", fill="both", expand=True)
        try:
            from utils import resource_path
            path = resource_path("logo.png")
            img = ctk.CTkImage(Image.open(path), size=(350, 350))
            ctk.CTkLabel(logo_fr, image=img, text="").place(relx=0.5, rely=0.5, anchor="center")
        except:
            ctk.CTkLabel(logo_fr, text="( LOGOTIP )", font=("Arial", 25), text_color="#333").place(relx=0.5, rely=0.5, anchor="center")

    def create_stat_box(self, parent, text, side):
        box = ctk.CTkFrame(parent, fg_color="transparent", border_width=1, border_color=SARIQ_RANG)
        box.pack(side=side, expand=True, fill="both", padx=20)
        lbl = ctk.CTkLabel(box, text=text, font=("Arial", 22, "bold"), text_color="white")
        lbl.place(relx=0.5, rely=0.5, anchor="center")
        return lbl

    def update_stats(self):
        self.lbl_debtors.configure(text=f"Qarzdorlar soni: {self.db.get_debtors_count()} ta")
        self.lbl_total.configure(text=f"Umumiy qarz: {format_money(self.db.get_total_debt_sum())}")

    # ================= 1-YANGI OYNA: MIJOZLAR MENYUSI (SIZNING RASMINGIZ) =================
    def open_customers_menu(self):
        m_win = ctk.CTkToplevel(self)
        m_win.title("Mijozlar menyusi")
        m_win.geometry("1100x700")
        m_win.configure(fg_color=FON_RANGI)
        m_win.grab_set()

        ctk.CTkLabel(m_win, text="SHOHJAXON QURILISH MARKET", font=("Oswald", 35, "bold"), text_color=SARIQ_RANG).pack(pady=25)
        
        cont = ctk.CTkFrame(m_win, fg_color="transparent")
        cont.pack(fill="both", expand=True, padx=40)

        side_menu = ctk.CTkFrame(cont, fg_color="transparent")
        side_menu.pack(side="left", fill="y", pady=10)

        # "mijoz qo'shish" tugmasi jadvalli oynani ochadi
        ctk.CTkButton(side_menu, text="mijoz qo'shish", width=280, height=60, font=("Arial", 17),
                      fg_color="transparent", border_width=2, border_color=SARIQ_RANG, text_color=SARIQ_RANG,
                      command=self.open_add_customer_window).pack(pady=15)
        
        ctk.CTkButton(side_menu, text="malumotlarni tahrirlash", width=280, height=60, 
                      fg_color="transparent", border_width=2, border_color=SARIQ_RANG, 
                      text_color=SARIQ_RANG, command=self.open_edit_customer_window).pack(pady=15)
         # 88-qator atrofida: "mijozni o'chirish" tugmasi
        ctk.CTkButton(side_menu, text="mijozni o'chirish", width=280, height=60, font=("Arial", 17),
                      fg_color="transparent", border_width=2, border_color=SARIQ_RANG, text_color=SARIQ_RANG,
                      command=self.open_delete_customer_window).pack(pady=15)

        self.add_logo(cont)

    # ================= 2-YANGI OYNA: MIJOZ QO'SHISH VA JADVAL =================
    def open_add_customer_window(self):
        a_win = ctk.CTkToplevel(self)
        a_win.title("Mijoz qo'shish")
        a_win.geometry("1100x750")
        a_win.configure(fg_color=FON_RANGI)
        a_win.grab_set()

        ctk.CTkLabel(a_win, text="YANGI MIJOZ QO'SHISH", font=("Oswald", 30, "bold"), text_color=SARIQ_RANG).pack(pady=20)

        f = ctk.CTkFrame(a_win, fg_color="transparent")
        f.pack(fill="x", padx=30)
        e1 = ctk.CTkEntry(f, placeholder_text="ism", width=220, height=40)
        e1.grid(row=0, column=0, padx=5)
        e2 = ctk.CTkEntry(f, placeholder_text="tel raqami", width=220, height=40)
        e2.grid(row=0, column=1, padx=5)
        e3 = ctk.CTkEntry(f, placeholder_text="manzil", width=220, height=40)
        e3.grid(row=0, column=2, padx=5)

        t_view = ctk.CTkTextbox(a_win, font=("Courier New", 14), fg_color="#0c0c0c", state="disabled", height=450)
        t_view.pack(padx=30, pady=20, fill="both")

        def ref():
            t_view.configure(state="normal"); t_view.delete("1.0", "end")
            header = f"{'№':<3} | {'ISM':<20} | {'TEL RAQAMI':<15} | {'MANZIL':<15} | {'QARZ'}\n"
            t_view.insert("end", header + "-"*75 + "\n")
            for i, c in enumerate(self.db.get_all_customers(), 1):
                debt = self.db.get_customer_total_debt(c.id)
                t_view.insert("end", f"{i:<3} | {str(c.full_name)[:19]:<20} | {str(c.phone)[:14]:<15} | {str(c.address)[:14]:<15} | {debt:,.0f}\n")
            t_view.configure(state="disabled")

        def save():
            if e1.get() and e2.get():
                self.db.add_customer(Customer(full_name=e1.get(), phone=e2.get(), address=e3.get()))
                e1.delete(0, 'end')
                e2.delete(0, 'end')
                e3.delete(0, 'end')
                
                # JADVALNI VA STATISTIKANI YANGILASH (SHU QATORLARNI QO'SHING)
                ref() 
                self.update_stats()
                
                messagebox.showinfo("OK", "Mijoz qo'shildi!")
            else:
                messagebox.showwarning("Xato", "Maydonlarni to'ldiring!")

        # Funksiyaning eng oxirgi qatori bo'lsin:
        ref()

        ctk.CTkButton(f, text="qo'shish", command=save, fg_color=SARIQ_RANG, text_color="black", font=("Arial", 14, "bold")).grid(row=0, column=3, padx=10)
    def open_edit_customer_window(self):
        from tkinter import ttk 
        e_win = ctk.CTkToplevel(self)
        e_win.title("Tahrirlash")
        e_win.geometry("1100x750")
        e_win.configure(fg_color=FON_RANGI)
        e_win.grab_set()

        ctk.CTkLabel(e_win, text="SHOHJAXON QURILISH MARKET", font=("Oswald", 35, "bold"), text_color=SARIQ_RANG).pack(pady=20)

        f = ctk.CTkFrame(e_win, fg_color="transparent")
        f.pack(fill="x", padx=30)
        
        # Kataklar
        e_id = ctk.CTkEntry(f, placeholder_text="ID", width=70, height=40)
        e_id.grid(row=0, column=0, padx=5)
        e_name = ctk.CTkEntry(f, placeholder_text="ism", width=200, height=40)
        e_name.grid(row=0, column=1, padx=5)
        e_phone = ctk.CTkEntry(f, placeholder_text="tel", width=200, height=40)
        e_phone.grid(row=0, column=2, padx=5)
        e_addr = ctk.CTkEntry(f, placeholder_text="manzil", width=200, height=40)
        e_addr.grid(row=0, column=3, padx=5)

        # Jadval dizayni
        style = ttk.Style()
        style.theme_use("default")
        style.configure("Treeview", background="#0c0c0c", foreground="white", fieldbackground="#0c0c0c", rowheight=35)
        style.map("Treeview", background=[('selected', SARIQ_RANG)], foreground=[('selected', 'black')])

        # Jadvalni yaratish
        tree = ttk.Treeview(e_win, columns=("ID", "Ism", "Tel", "Manzil"), show='headings')
        tree.heading("ID", text="ID"); tree.column("ID", width=50)
        tree.heading("Ism", text="MIJOZ ISMI"); tree.column("Ism", width=250)
        tree.heading("Tel", text="TELEFON"); tree.column("Tel", width=200)
        tree.heading("Manzil", text="MANZIL"); tree.column("Manzil", width=250)
        tree.pack(padx=30, pady=20, fill="both", expand=True)

        def ref():
            for i in tree.get_children(): tree.delete(i)
            for c in self.db.get_all_customers():
                tree.insert("", "end", values=(c.id, c.full_name, c.phone, c.address))

        # Mijoz ustiga bosilganda kataklarni to'ldirish
        def on_select(event):
            selected = tree.selection()
            if selected:
                val = tree.item(selected)['values']
                e_id.delete(0, 'end'); e_id.insert(0, val[0])
                e_name.delete(0, 'end'); e_name.insert(0, val[1])
                e_phone.delete(0, 'end'); e_phone.insert(0, val[2])
                e_addr.delete(0, 'end'); e_addr.insert(0, val[3])

        tree.bind("<<TreeviewSelect>>", on_select)

        def update():
            if e_id.get():
                self.db.update_customer(e_id.get(), e_name.get(), e_phone.get(), e_addr.get())
                messagebox.showinfo("OK", "Ma'lumotlar tahrirlandi!")
                ref(); self.update_stats()
            else: messagebox.showwarning("Xato", "Mijozni tanlang!")

        ctk.CTkButton(f, text="tahrirlash", command=update, fg_color=SARIQ_RANG, text_color="black", font=("Arial", 14, "bold")).grid(row=0, column=4, padx=10)
        ref()
    def open_delete_customer_window(self):
        from tkinter import ttk, messagebox
        d_win = ctk.CTkToplevel(self)
        d_win.title("Mijozni o'chirish")
        d_win.geometry("1100x750")
        d_win.configure(fg_color=FON_RANGI)
        d_win.grab_set()

        # 1. SARLAVHA
        ctk.CTkLabel(d_win, text="SHOHJAXON QURILISH MARKET", 
                     font=("Oswald", 35, "bold"), text_color=SARIQ_RANG).pack(pady=20)

        # 2. QIDIRUV VA TUGMA KONTEYNERI
        mid_fr = ctk.CTkFrame(d_win, fg_color="transparent")
        mid_fr.pack(fill="x", padx=30, pady=10)

        search_entry = ctk.CTkEntry(mid_fr, placeholder_text="ism yozing...", width=300, height=45)
        search_entry.pack(side="left", padx=10)

        # 3. JADVAL (TREEVIEW)
        style = ttk.Style()
        style.theme_use("default")
        style.configure("Treeview", background="#0c0c0c", foreground="white", fieldbackground="#0c0c0c", rowheight=35)
        style.map("Treeview", background=[('selected', '#e74c3c')], foreground=[('selected', 'white')])

        tree = ttk.Treeview(d_win, columns=("№", "Ism", "Tel", "Manzil", "Qarz"), show='headings')
        tree.heading("№", text="№"); tree.column("№", width=50)
        tree.heading("Ism", text="ISM"); tree.column("Ism", width=250)
        tree.heading("Tel", text="TEL RAQAMI"); tree.column("Tel", width=180)
        tree.heading("Manzil", text="MANZIL"); tree.column("Manzil", width=220)
        tree.heading("Qarz", text="UMUMIY QARZ MIQDORI"); tree.column("Qarz", width=200)
        tree.pack(padx=30, pady=20, fill="both", expand=True)

        def ref(search_query=""):
            for i in tree.get_children(): tree.delete(i)
            customers = self.db.get_all_customers()
            filtered = [c for c in customers if search_query.lower() in c.full_name.lower()]
            for i, c in enumerate(filtered, 1):
                debt = self.db.get_customer_total_debt(c.id)
                tree.insert("", "end", values=(i, c.full_name, c.phone, c.address, f"{debt:,.0f} so'm"), tags=(c.id,))

        search_entry.bind("<KeyRelease>", lambda e: ref(search_entry.get()))

        def delete_now():
            selected = tree.selection()
            if selected:
                item = tree.item(selected)
                c_id = item['tags'][0]
                c_name = item['values'][1]
                if messagebox.askyesno("Tasdiqlash", f"Diqqat! {c_name}ni o'chirmoqchimisiz?"):
                    self.db.delete_customer(c_id)
                    ref(search_entry.get())
                    self.update_stats()
                    messagebox.showinfo("OK", "Mijoz o'chirildi.")
            else:
                messagebox.showwarning("Xato", "O'chirish uchun mijozni tanlang!")

        # O'CHIRISH TUGMASI
        ctk.CTkButton(mid_fr, text="o'chirish", width=180, height=45, 
                      fg_color="#e74c3c", text_color="white", font=("Arial", 15, "bold"),
                      command=delete_now).pack(side="right", padx=10)

        ref()
    def open_add_debt_window(self):
        from tkinter import ttk, messagebox
        ad_win = ctk.CTkToplevel(self)
        ad_win.title("Qarz qo'shish")
        ad_win.geometry("1150x750")
        ad_win.configure(fg_color=FON_RANGI)
        ad_win.grab_set()

        ctk.CTkLabel(ad_win, text="SHOHJAXON QURILISH MARKET", font=("Oswald", 35, "bold"), text_color=SARIQ_RANG).pack(pady=20)

        f = ctk.CTkFrame(ad_win, fg_color="transparent")
        f.pack(fill="x", padx=30, pady=10)
        
        e_name = ctk.CTkEntry(f, placeholder_text="ism", width=220, height=45)
        e_name.grid(row=0, column=0, padx=5)
        e_amount = ctk.CTkEntry(f, placeholder_text="qarz miqdori", width=220, height=45)
        e_amount.grid(row=0, column=1, padx=5)
        e_desc = ctk.CTkEntry(f, placeholder_text="izoh (tovar nomi)", width=220, height=45)
        e_desc.grid(row=0, column=2, padx=5)

        self.selected_cust_id = None 

        style = ttk.Style()
        style.configure("Treeview", background="#0c0c0c", foreground="white", fieldbackground="#0c0c0c", rowheight=35)
        style.map("Treeview", background=[('selected', SARIQ_RANG)], foreground=[('selected', 'black')])

        tree = ttk.Treeview(ad_win, columns=("№", "Ism", "Tel", "Manzil", "Qarz"), show='headings')
        tree.heading("№", text="№"); tree.column("№", width=50)
        tree.heading("Ism", text="ISM"); tree.column("Ism", width=250)
        tree.heading("Tel", text="TEL RAQAMI"); tree.column("Tel", width=180)
        tree.heading("Manzil", text="MANZIL"); tree.column("Manzil", width=220)
        tree.heading("Qarz", text="UMUMIY QARZ"); tree.column("Qarz", width=200)
        tree.pack(padx=30, pady=20, fill="both", expand=True)

        def ref(q=""):
            for i in tree.get_children(): tree.delete(i)
            custs = [c for c in self.db.get_all_customers() if q.lower() in c.full_name.lower()]
            for i, c in enumerate(custs, 1):
                debt = self.db.get_customer_total_debt(c.id)
                tree.insert("", "end", values=(i, c.full_name, c.phone, c.address, f"{debt:,.0f}"), tags=(c.id,))

        e_name.bind("<KeyRelease>", lambda e: ref(e_name.get()))

        def on_select(event):
            sel = tree.selection()
            if sel:
                val = tree.item(sel)['values']
                self.selected_cust_id = tree.item(sel)['tags'][0]
                e_name.delete(0, 'end'); e_name.insert(0, val[1]) # To'liq ismni yozadi

        tree.bind("<<TreeviewSelect>>", on_select)

        def save_debt():
            summa = e_amount.get().strip()
            if self.selected_cust_id and summa:
                try:
                    amount = float(summa)
                    self.db.add_debt(self.selected_cust_id, amount, e_desc.get())
                    messagebox.showinfo("OK", "Qarz muvaffaqiyatli qo'shildi!")
                    e_amount.delete(0, 'end')
                    e_desc.delete(0, 'end')
                    ref()
                    self.update_stats()
                except ValueError:
                    messagebox.showerror("Xato", "Faqat raqam kiriting!")
            else:
                messagebox.showwarning("Xato", "Mijozni tanlang va summani yozing!")

        ctk.CTkButton(f, text="qo'shish", width=180, height=45, 
                      fg_color=SARIQ_RANG, text_color="black", 
                      font=("Arial", 15, "bold"), command=save_debt).grid(row=0, column=3, padx=10)
        ref()
    def open_payment_window(self):
        from tkinter import ttk, messagebox
        p_win = ctk.CTkToplevel(self)
        p_win.title("To'lov qabul qilish")
        p_win.geometry("1150x750")
        p_win.configure(fg_color=FON_RANGI)
        p_win.grab_set()

        # 1. DO'KON NOMI (YUQORI QISM)
        ctk.CTkLabel(p_win, text="SHOHJAXON QURILISH MARKET", 
                     font=("Oswald", 35, "bold"), text_color=SARIQ_RANG).pack(pady=20)

        # 2. KIRITISH MAYDONLARI (CHIZMADAGIDEK)
        f = ctk.CTkFrame(p_win, fg_color="transparent")
        f.pack(fill="x", padx=30, pady=10)
        
        e_name = ctk.CTkEntry(f, placeholder_text="ism", width=220, height=45)
        e_name.grid(row=0, column=0, padx=5)
        e_pay = ctk.CTkEntry(f, placeholder_text="to'lov miqdori", width=220, height=45)
        e_pay.grid(row=0, column=1, padx=5)
        e_desc = ctk.CTkEntry(f, placeholder_text="izoh", width=220, height=45)
        e_desc.grid(row=0, column=2, padx=5)

        self.pay_cust_id = None 

        # 3. JADVAL (MIJOZLAR RO'YXATI)
        style = ttk.Style()
        style.configure("Treeview", background="#0c0c0c", foreground="white", fieldbackground="#0c0c0c", rowheight=35)
        style.map("Treeview", background=[('selected', "#27ae60")], foreground=[('selected', 'white')])

        tree = ttk.Treeview(p_win, columns=("№", "Ism", "Tel", "Manzil", "Qarz"), show='headings')
        tree.heading("№", text="№"); tree.column("№", width=50)
        tree.heading("Ism", text="ISM"); tree.column("Ism", width=250)
        tree.heading("Tel", text="TEL RAQAMI"); tree.column("Tel", width=180)
        tree.heading("Manzil", text="MANZIL"); tree.column("Manzil", width=220)
        tree.heading("Qarz", text="UMUMIY QARZ MIQDORI"); tree.column("Qarz", width=200)
        tree.pack(padx=30, pady=20, fill="both", expand=True)

        def ref(q=""):
            for i in tree.get_children(): tree.delete(i)
            custs = [c for c in self.db.get_all_customers() if q.lower() in c.full_name.lower()]
            for i, c in enumerate(custs, 1):
                debt = self.db.get_customer_total_debt(c.id)
                tree.insert("", "end", values=(i, c.full_name, c.phone, c.address, f"{debt:,.0f} so'm"), tags=(c.id,))

        e_name.bind("<KeyRelease>", lambda e: ref(e_name.get()))

        def on_select(event):
            sel = tree.selection()
            if sel:
                val = tree.item(sel)['values']
                # List xatosi bo'lmasligi uchun ID ni tags'dan to'g'ri olamiz
                self.pay_cust_id = tree.item(sel)['tags'][0]
                e_name.delete(0, 'end'); e_name.insert(0, val[1])

        tree.bind("<<TreeviewSelect>>", on_select)

        def make_payment():
            summa_txt = e_pay.get().strip()
            if self.pay_cust_id and summa_txt:
                try:
                    amount = float(summa_txt)
                    izoh = e_desc.get().strip() if e_desc.get() else "To'lov qabul qilindi"
                    self.db.add_payment(self.pay_cust_id, amount, izoh)
                    messagebox.showinfo("OK", f"{amount:,.0f} so'm to'lov qabul qilindi!")
                    e_pay.delete(0, 'end'); e_desc.delete(0, 'end'); ref(); self.update_stats()
                except:
                    messagebox.showerror("Xato", "To'lov miqdorini raqamda kiriting!")
            else:
                messagebox.showwarning("Xato", "Mijozni tanlang va summani yozing!")

        # 4. TO'LOV TUGMASI (CHIZMADAGIDEK O'NG TOMONDA)
        ctk.CTkButton(f, text="to'lov tugmasi", width=180, height=45, 
                      fg_color="#27ae60", text_color="white", font=("Arial", 15, "bold"),
                      command=make_payment).grid(row=0, column=3, padx=10)
        ref()
    def open_calculations_window(self):
        calc_win = ctk.CTkToplevel(self)
        calc_win.title("Hisob-kitoblar")
        calc_win.geometry("1100x600")
        calc_win.configure(fg_color=FON_RANGI)
        calc_win.grab_set()

        # 1. DO'KON NOMI
        ctk.CTkLabel(calc_win, text="SHOHJAXON QURILISH MARKET", font=("Oswald", 30, "bold"), text_color=SARIQ_RANG).pack(pady=2)

        # 2. QIDIRUV BO'LIMI
        search_fr = ctk.CTkFrame(calc_win, fg_color="transparent")
        search_fr.pack(fill="x", padx=30, pady=2)
        
        e_search = ctk.CTkEntry(search_fr, placeholder_text="qidiruv uchun joy", width=250)
        e_search.pack(side="left", padx=5)
        
        ctk.CTkButton(search_fr, text="🔍 qidiruv", width=100, fg_color=SARIQ_RANG, text_color="black", 
                      command=lambda: search_customer(e_search.get())).pack(side="left")

        # 3. YUQORI JADVAL
        from tkinter import ttk
        columns = ("no", "ism", "tel", "manzil", "qarz")
        tree1 = ttk.Treeview(calc_win, columns=columns, show="headings", height=4)
        for col in columns: tree1.heading(col, text=col.upper())
        tree1.pack(fill="both", expand=True, padx=30, pady=10)

        # IKKI MARTA BOSISH (DOUBLE CLICK)
        tree1.bind("<Double-1>", lambda event: self.open_details_window(tree1.item(tree1.selection(), "values")))

        # 4. FILTR BO'LIMI
        filter_fr = ctk.CTkFrame(calc_win, fg_color="transparent")
        filter_fr.pack(fill="x", padx=30, pady=5)
        ctk.CTkButton(filter_fr, text="📑 filtr", width=100, border_width=1, 
                      border_color=SARIQ_RANG, fg_color="transparent", text_color=SARIQ_RANG,
                      command=self.open_filter_window).pack(side="left", padx=10)

        # 5. PASTORGI JADVAL
        self.tree2 = ttk.Treeview(calc_win, columns=columns, show="headings", height=4)
        for col in columns: self.tree2.heading(col, text=col.upper())
        self.tree2.pack(fill="both", expand=True, padx=30, pady=10)

        # 6. CHOP QILISH
        ctk.CTkButton(calc_win, text="chop qilish", fg_color="white", text_color="black", 
                      font=("Arial", 14, "bold"), width=150, height=40,
                      command=self.export_filtered_data).pack(side="right", padx=30, pady=10)

        def search_customer(text=""):
            for i in tree1.get_children(): tree1.delete(i)
            for i, c in enumerate(self.db.get_all_customers(), 1):
                if text.lower() in c.full_name.lower():
                    debt = self.db.get_customer_total_debt(c.id)
                    tree1.insert("", "end", values=(i, c.full_name, c.phone, c.address, f"{debt:,.0f}"))
        search_customer()

    def open_details_window(self, data):
        if not data: return
        
        # data ichidan ma'lumotlarni ajratib olamiz
        c_id, c_name, c_phone, c_addr, c_total_debt = data[0], data[1], data[2], data[3], data[4]
        
        det_win = ctk.CTkToplevel(self)
        det_win.title(f"{c_name} - Batafsil ma'lumot")
        det_win.geometry("1000x850")
        det_win.configure(fg_color=FON_RANGI)
        det_win.grab_set()

        # 1. DO'KON NOMI
        ctk.CTkLabel(det_win, text="SHOHJAXON QURILISH MARKET", font=("Oswald", 30, "bold"), text_color=SARIQ_RANG).pack(pady=15)

        # 2. MIJOZ MA'LUMOTLARI (SHABLONDAGI 2-QATOR)
        info_fr = ctk.CTkFrame(det_win, fg_color="transparent", border_width=1, border_color=SARIQ_RANG)
        info_fr.pack(fill="x", padx=30, pady=10)
        
        headers = ["ism", "tel raqami", "manzil", "umumiy qarz miqdori"]
        values = [c_name, c_phone, c_addr, c_total_debt]
        
        for i in range(4):
            f = ctk.CTkFrame(info_fr, fg_color="transparent", border_width=1, border_color="#333")
            f.grid(row=0, column=i, sticky="nsew")
            info_fr.grid_columnconfigure(i, weight=1)
            ctk.CTkLabel(f, text=headers[i], font=("Arial", 13, "bold"), text_color=SARIQ_RANG).pack(pady=2)
            ctk.CTkLabel(f, text=values[i], font=("Arial", 15), text_color="white").pack(pady=5)

        # 3. OPERATSIYALAR JADVALI (№, Miqdor, Sana, Izoh)
        from tkinter import ttk
        columns = ("no", "miqdor", "sana", "izoh")
        tree_det = ttk.Treeview(det_win, columns=columns, show="headings", height=15)
        
        tree_det.heading("no", text="№")
        tree_det.heading("miqdor", text="olingan qarz yoki qilingan to'lov miqdori")
        tree_det.heading("sana", text="sana")
        tree_det.heading("izoh", text="izoh")
        
        # Ustunlar kengligini sozlash
        tree_det.column("no", width=50, anchor="center")
        tree_det.column("miqdor", width=300, anchor="center")
        tree_det.column("sana", width=150, anchor="center")
        tree_det.column("izoh", width=300, anchor="w")
        
        tree_det.pack(fill="both", expand=True, padx=30, pady=10)
        # Jadvalni to'ldirish
        history = self.db.get_customer_history(data[0]) # data[0] - bu mijoz ID si
        for i, row in enumerate(history, 1):
            amount, date, comment = row
            # To'lov bo'lsa (manfiy summa) ajralib turishi uchun formatlash
            m_text = f"{amount:,.0f}"
            tree_det.insert("", "end", values=(i, m_text, date, comment))

        # 4. UMUMIY QARZ PASTI (SHABLONDAGI KABI)
        bottom_fr = ctk.CTkFrame(det_win, fg_color="transparent")
        bottom_fr.pack(fill="x", padx=30)
        ctk.CTkLabel(bottom_fr, text=f"umumiy qarz: {c_total_debt}", font=("Arial", 18, "bold"), text_color=SARIQ_RANG).pack(side="left", padx=50)

        # 5. CHOP QILISH TUGMASI (O'NG PASTDA)
        ctk.CTkButton(det_win, text="chop qilish", fg_color="white", text_color="black", 
                      font=("Arial", 14, "bold"), width=150, height=45,
                      command=lambda: self.export_to_excel(c_name, history)).pack(side="right", padx=30, pady=20)
    def export_to_excel(self, name, history):
        import pandas as pd
        from tkinter import filedialog, messagebox

        try:
            path = filedialog.asksaveasfilename(
                defaultextension=".xlsx", 
                initialfile=f"{name}_tarixi.xlsx",
                title="Hisobotni saqlash",
                filetypes=[
                    ("Excel fayllari", "*.xlsx"),
                    ("Word hujjatlari", "*.docx"),
                    ("PDF hujjatlari", "*.pdf"),
                    ("Barcha fayllar", "*.*")
                ]
            )
            
            if path:
                if path.endswith(".xlsx"):
                    df = pd.DataFrame(history, columns=["Miqdor", "Sana", "Izoh"])
                    df.to_excel(path, index=False)
                elif path.endswith(".docx"):
                    self.save_as_word(path, name, history)
                elif path.endswith(".pdf"):
                    self.save_as_pdf(path, name, history)
                
                messagebox.showinfo("Tayyor", f"{name} uchun hisobot saqlandi!")

        except Exception as e:
            messagebox.showerror("Xato", f"Saqlashda xatolik yuz berdi: {e}")
            from docx import Document # Kutubxonani funksiya ichida chaqiramiz
    def save_as_word(self, path, name, history):
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(f"{name}", 0)
            
            # Agar history bo'sh bo'lmasa, birinchi qatordagi ustunlar sonini olamiz
            cols_count = len(history[0]) if history else 3
            table = doc.add_table(rows=1, cols=cols_count)
            table.style = 'Table Grid'
            
            # Sarlavhalarni ustunlar soniga qarab belgilash
            hdr_cells = table.rows[0].cells
            if cols_count == 5:
                labels = ["№", "Ism", "Tel", "Manzil", "Qarz"]
            else:
                labels = ["Miqdor", "Sana", "Izoh"]
                
            for i in range(min(cols_count, len(labels))):
                hdr_cells[i].text = labels[i]
            
            # Ma'lumotlarni qatorlar bo'yicha qo'shish
            for row in history:
                row_cells = table.add_row().cells
                for i in range(len(row)):
                    row_cells[i].text = str(row[i])
                    
            doc.save(path)
        except Exception as e:
            print(f"Word xatosi: {e}")
    def save_as_pdf(self, path, name, history):
        from fpdf import FPDF
        try:
            pdf = FPDF()
            pdf.add_page()
            # Standart shrift (O'zbekcha harflar uchun keyinroq font qo'shamiz)
            pdf.set_font("Arial", size=10)
            pdf.cell(200, 10, txt=f"{name}", ln=1, align='C')
            
            for row in history:
                # 3 ta yoki 5 ta ustun bo'lishidan qat'i nazar hammasini chiqaradi
                line = " | ".join(map(str, row))
                pdf.cell(200, 10, txt=line, ln=1)
            pdf.output(path)
        except Exception as e:
            print(f"PDF xatosi: {e}")

    def save_as_pdf(self, path, name, history):
        from fpdf import FPDF
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=10)
            
            pdf.cell(200, 10, txt=str(name), ln=1, align='C')
            pdf.ln(5)
            
            for row in history:
                # 3 ta yoki 5 ta ustun bo'lishidan qat'i nazar bitta qatorga yozadi
                line = " | ".join(map(str, row))
                pdf.cell(0, 10, txt=line, ln=1)
                
            pdf.output(path)
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror("PDF xatosi", f"Xato yuz berdi: {e}")
    def open_filter_window(self):
        f_win = ctk.CTkToplevel(self)
        f_win.title("Filtrlash")
        f_win.geometry("450x420")
        f_win.configure(fg_color=FON_RANGI)
        f_win.grab_set()

        ctk.CTkLabel(f_win, text="FILTRLASH", font=("Oswald", 25, "bold"), text_color=SARIQ_RANG).pack(pady=15)

        # Miqdorlar (Yonma-yon)
        am_fr = ctk.CTkFrame(f_win, fg_color="transparent")
        am_fr.pack(fill="x", padx=25, pady=10)
        e_min = ctk.CTkEntry(am_fr, placeholder_text="eng kichik miqdor", width=190)
        e_min.pack(side="left", padx=5)
        e_max = ctk.CTkEntry(am_fr, placeholder_text="eng katta miqdor", width=190)
        e_max.pack(side="right", padx=5)

        # Manzil va Sana
        e_addr = ctk.CTkEntry(f_win, placeholder_text="manzil", width=390)
        e_addr.pack(pady=10)
        e_date = ctk.CTkEntry(f_win, placeholder_text="sana (kk.oo.yyyy)", width=390)
        e_date.pack(pady=10)

        # Saralash tugmasi
        def apply_filter():
            # 1. Маълумотларни олиш
            mi_d = e_min.get()
            ma_d = e_max.get()
            addr = e_addr.get()
            
            # 2. Пастки жадвални (tree2) тозалаш
            for i in self.tree2.get_children():
                self.tree2.delete(i)
                
            # 3. Базадан қидириш
            res = self.db.get_filtered_customers(
                min_debt=mi_d if mi_d else None,
                max_debt=ma_d if ma_d else None,
                address=addr if addr else None
            )
            
            # 4. Жадвалга чиқариш
            for i, (c, debt) in enumerate(res, 1):
                # c - бу мижоз объекти (full_name, phone, address каби хусусиятлари билан)
            # c[1] - исм, c[2] - телефон, c[3] - манзил
                self.tree2.insert("", "end", values=(i, c[1], c[2], c[3], f"{debt:,.0f}"))
          
            f_win.destroy()
        # Фақат мана шу битта тугма код блоги қолсин:
        # 676-qatordan boshlab hamma kodni mana bunga almashtiring:
        ctk.CTkButton(f_win, text="saralash", fg_color=SARIQ_RANG, text_color="black",
                      font=("Arial", 16, "bold"), width=250, height=45,
                      command=apply_filter).pack(pady=25)
    def export_filtered_data(self):
        import pandas as pd
        from tkinter import filedialog, messagebox

        # Агар пастки жадвал бўш бўлса
        if not self.tree2.get_children():
            messagebox.showwarning("Диққат", "Аввал саралаш тугмасини босиб маълумотларни чиқаринг!")
            return

        # 692-қатордан бошлаб текширинг:
        path = filedialog.asksaveasfilename(
            defaultextension=".xlsx", 
            initialfile="saralangan_mijozlar.xlsx",
            title="Ҳисоботни сақлаш",
            filetypes=[
                ("Excel fayllari", "*.xlsx"),
                ("Word hujjatlari", "*.docx"),
                ("PDF hujjatlari", "*.pdf"),
                ("Barcha fayllar", "*.*")
            ]
        )
        
        if path: # Энди 'path' юқорида эълон қилинди, хато бермайди
            try:
                data = []
                for item in self.tree2.get_children():
                    data.append(self.tree2.item(item)["values"])
                
                doc_title = "Saralangan Mijozlar Ro'yxati"
                cols = ["№", "Ism", "Tel", "Manzil", "Qarz"]

                if path.endswith(".xlsx"):
                    df = pd.DataFrame(data, columns=cols)
                    df.to_excel(path, index=False)
                elif path.endswith(".docx"):
                    self.save_as_word(path, doc_title, data)
                elif path.endswith(".pdf"):
                    self.save_as_pdf(path, doc_title, data)
                
                messagebox.showinfo("Tayyor", "Ma'lumotlar muvaffaqiyatli saqlandi!")

            except Exception as e:
                messagebox.showerror("Xato", f"Saqlashda xatolik: {e}")
class LoginWindow(ctk.CTk):
    def __init__(self, on_success):
        super().__init__()
        self.on_success = on_success
        self.db = Database()
        self.title("Kirish - SHOHJAXON MARKET")
        self.geometry("400x500")
        self.configure(fg_color="#0a0a0a")

        # Sarlavha
        ctk.CTkLabel(self, text="TIZIMGA KIRISH", font=("Oswald", 25, "bold"), text_color="#f1c40f").pack(pady=40)
        
        # Email kiritish joyi
        self.e_mail = ctk.CTkEntry(self, placeholder_text="Email", width=280, height=45)
        self.e_mail.pack(pady=10)
        
        # Parol kiritish joyi
        self.e_pass = ctk.CTkEntry(self, placeholder_text="Parol", width=280, height=45, show="*")
        self.e_pass.pack(pady=10)
        
        # Kirish tugmasi
        btn = ctk.CTkButton(self, text="KIRISH", fg_color="#f1c40f", text_color="black", 
                            font=("Arial", 16, "bold"), height=45, width=280, command=self.try_login)
        btn.pack(pady=30)

    def try_login(self):
        # Database.py dagi login funksiyasini chaqiramiz
        res = self.db.login(self.e_mail.get(), self.e_pass.get())
        
        if res == "SUCCESS":
            self.destroy()      # Login oynasini yopish
            self.on_success()   # Asosiy dasturni (MarketApp) ochish
        elif res == "CONFIRM_REQUIRED":
            messagebox.showwarning("Diqqat", "Iltimos, pochtangizga kelgan xatni tasdiqlang!")
        else:
            messagebox.showerror("Xato", "Email yoki parol noto'g'ri!")
